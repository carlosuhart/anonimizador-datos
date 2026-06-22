# =============================================================================
# test_anonimizar.py — Suite de regresión del anonimizador
#
# Ejecuta el script real (subprocess) sobre fixtures con PII conocida de cada
# jurisdicción y verifica que:
#   - cada identificador se tokeniza (no aparece literal en la salida)
#   - el round-trip anonimizar -> restaurar devuelve el original
#   - el cifrado del mapa protege la PII y la clave es obligatoria al restaurar
#
# Uso:  python -m unittest test_anonimizar       (requiere modelos de spaCy)
#       python test_anonimizar.py
#
# Nota: cada anonimización carga los modelos NLP (~10-20 s). La suite agrupa los
# casos para minimizar el número de cargas.
# =============================================================================

import os
import sys
import json
import shutil
import tempfile
import subprocess
import unittest

DIR = os.path.dirname(os.path.abspath(__file__))
SCRIPT = os.path.join(DIR, "anonimizar.py")

# Identificadores por jurisdicción: deben desaparecer del archivo anonimizado.
IDENTIFICADORES = {
    "DNI-ES":   "12345678Z",
    "RUT-CL":   "12.345.678-5",
    "CPF-BR":   "123.456.789-09",
    "CNPJ-BR":  "12.345.678/0001-95",
    "CURP-MX":  "PELJ800101HDFRRN09",
    "NIT-CO":   "123456789-1",
    "CUIT-AR":  "20-12345678-9",
    "NINO-UK":  "AB123456C",
    "SSN-US":   "123-45-6789",
    "EMAIL":    "juan.perez@empresa.com",
    "DIRECCION": "Calle Mayor 3",
}


def _run(args, cwd=DIR):
    # Fixtures en español: cargar solo el modelo es acelera cada subprocess ~3x.
    # stdin=DEVNULL: ejecución no interactiva (como la de un agente). Evita que
    # getpass quede esperando input al restaurar un mapa cifrado sin clave.
    env = dict(os.environ, ANON_IDIOMAS="es")
    env.pop("ANON_CLAVE", None)
    return subprocess.run(
        [sys.executable, SCRIPT, *args],
        cwd=cwd, capture_output=True, text=True, env=env,
        stdin=subprocess.DEVNULL,
    )


class TestAnonimizacion(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix="anon_test_")
        cls.csv = os.path.join(cls.tmp, "fixture.csv")
        filas = ["nombre,email,documento,direccion"]
        for i, (_, valor) in enumerate(IDENTIFICADORES.items(), 1):
            # Cada fila tiene un nombre (activa contexto) + un identificador.
            filas.append(f"Persona Ejemplo {i},juan.perez@empresa.com,{valor},Calle Mayor 3")
        with open(cls.csv, "w", encoding="utf-8") as f:
            f.write("\n".join(filas) + "\n")
        cls.res = _run([cls.csv, "--ley", "todo"])
        cls.anon = os.path.join(cls.tmp, "fixture_anon.csv")
        cls.key = cls.anon + ".key.json"

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_proceso_termina_ok(self):
        self.assertEqual(self.res.returncode, 0, self.res.stderr)
        self.assertTrue(os.path.isfile(self.anon))
        self.assertTrue(os.path.isfile(self.key))

    def test_identificadores_tokenizados(self):
        with open(self.anon, encoding="utf-8") as f:
            contenido = f.read()
        for etiqueta, valor in IDENTIFICADORES.items():
            with self.subTest(identificador=etiqueta):
                self.assertNotIn(valor, contenido,
                                 f"{etiqueta} ({valor}) NO se anonimizó — fuga de PII")

    def test_mapa_recupera_valores(self):
        with open(self.key, encoding="utf-8") as f:
            mapa = json.load(f)["mapa"]
        originales = set(mapa.values())
        for etiqueta, valor in IDENTIFICADORES.items():
            with self.subTest(identificador=etiqueta):
                self.assertIn(valor, originales, f"{etiqueta} no está en el mapa")

    def test_reporte_cobertura_en_salida(self):
        self.assertIn("Cobertura de detección", self.res.stdout)

    def test_roundtrip_restaura_original(self):
        r = _run([self.anon, "--restaurar"])
        self.assertEqual(r.returncode, 0, r.stderr)
        restaurado = os.path.join(self.tmp, "fixture_anon_restaurado.csv")
        with open(self.csv, encoding="utf-8") as f:
            original = f.read()
        with open(restaurado, encoding="utf-8") as f:
            self.assertEqual(f.read(), original)


class TestCifrado(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix="anon_cifr_")
        cls.csv = os.path.join(cls.tmp, "datos.csv")
        with open(cls.csv, "w", encoding="utf-8") as f:
            f.write("nombre,documento\nAna Ruiz,12345678Z\n")
        cls.res = _run([cls.csv, "--ley", "rgpd", "--cifrar-mapa", "--clave", "secreta123"])
        cls.anon = os.path.join(cls.tmp, "datos_anon.csv")
        cls.key = cls.anon + ".key.json"

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_mapa_cifrado_sin_pii_en_claro(self):
        self.assertEqual(self.res.returncode, 0, self.res.stderr)
        with open(self.key, encoding="utf-8") as f:
            sobre = json.load(f)
        self.assertTrue(sobre.get("cifrado"))
        self.assertNotIn("12345678Z", json.dumps(sobre))

    def test_restaurar_sin_clave_falla(self):
        # Aislar de otros tests: borrar cualquier restaurado previo
        restaurado = os.path.join(self.tmp, "datos_anon_restaurado.csv")
        if os.path.isfile(restaurado):
            os.remove(restaurado)
        r = _run([self.anon, "--restaurar"])
        self.assertIn("cifrado", (r.stdout + r.stderr).lower())
        self.assertFalse(os.path.isfile(restaurado),
                         "No debe generarse el restaurado sin la clave")

    def test_restaurar_con_clave_ok(self):
        r = _run([self.anon, "--restaurar", "--clave", "secreta123"])
        self.assertEqual(r.returncode, 0, r.stderr)
        restaurado = os.path.join(self.tmp, "datos_anon_restaurado.csv")
        with open(restaurado, encoding="utf-8") as f:
            self.assertIn("12345678Z", f.read())


class TestDeteccionYFormatos(unittest.TestCase):
    """Cubre lo frágil: NER de nombres, formatos .md/.docx end-to-end,
    contexto de ubicación por fila y ausencia de falsos positivos. Todos los
    archivos se anonimizan en UNA invocación (--ley rgpd) para una sola carga."""

    @classmethod
    def setUpClass(cls):
        from docx import Document
        cls.tmp = tempfile.mkdtemp(prefix="anon_cob_")
        p = lambda n: os.path.join(cls.tmp, n)

        # .md con nombre (NER) + dirección
        cls.md = p("nota.md")
        with open(cls.md, "w", encoding="utf-8") as f:
            f.write("Cita con Juan García en la Calle Mayor 3 el lunes.\n")

        # .docx con nombre (NER) + DNI, nombre partido en dos runs
        cls.docx = p("informe.docx")
        doc = Document()
        par = doc.add_paragraph()
        par.add_run("Paciente: María ")
        par.add_run("López, DNI 12345678Z")
        doc.save(cls.docx)

        # CSV con persona + ciudad (la ciudad debe anonimizarse por contexto de fila)
        cls.csv_pers = p("personas.csv")
        with open(cls.csv_pers, "w", encoding="utf-8") as f:
            f.write("nombre,ciudad\nAna Ruiz,Sevilla\n")

        # CSV sin personas: la ciudad NO debe tocarse (evita falso positivo)
        cls.csv_inv = p("inventario.csv")
        with open(cls.csv_inv, "w", encoding="utf-8") as f:
            f.write("producto,region\nTeclado,Sevilla\n")

        # CSV con un código numérico no personal: no debe tokenizarse con rgpd
        cls.csv_sku = p("pedidos.csv")
        with open(cls.csv_sku, "w", encoding="utf-8") as f:
            f.write("detalle\nPedido 4521 procesado\n")

        cls.res = _run([cls.md, cls.docx, cls.csv_pers, cls.csv_inv, cls.csv_sku, "--ley", "rgpd"])

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def _leer(self, ruta):
        with open(ruta, encoding="utf-8") as f:
            return f.read()

    def test_proceso_ok(self):
        self.assertEqual(self.res.returncode, 0, self.res.stderr)

    def test_md_nombre_y_direccion(self):
        out = self._leer(os.path.join(self.tmp, "nota_anon.md"))
        self.assertNotIn("Juan García", out, "nombre no detectado en .md")
        self.assertNotIn("Calle Mayor 3", out, "dirección no detectada en .md")

    def test_docx_nombre_partido_y_dni(self):
        from docx import Document
        out = " ".join(p.text for p in Document(os.path.join(self.tmp, "informe_anon.docx")).paragraphs)
        self.assertNotIn("María López", out, "nombre partido en runs no detectado en .docx")
        self.assertNotIn("12345678Z", out, "DNI no detectado en .docx")

    def test_ciudad_con_persona_se_anonimiza(self):
        out = self._leer(os.path.join(self.tmp, "personas_anon.csv"))
        self.assertNotIn("Sevilla", out, "ciudad junto a persona no anonimizada (contexto de fila)")

    def test_ciudad_sin_persona_intacta(self):
        out = self._leer(os.path.join(self.tmp, "inventario_anon.csv"))
        self.assertIn("Sevilla", out, "ciudad sin persona NO debería anonimizarse")

    def test_codigo_no_es_falso_positivo(self):
        out = self._leer(os.path.join(self.tmp, "pedidos_anon.csv"))
        self.assertIn("4521", out, "un código numérico no personal no debe tokenizarse con rgpd")

    def test_roundtrip_md(self):
        r = _run([os.path.join(self.tmp, "nota_anon.md"), "--restaurar"])
        self.assertEqual(r.returncode, 0, r.stderr)
        self.assertEqual(self._leer(os.path.join(self.tmp, "nota_anon_restaurado.md")),
                         self._leer(self.md))

    def test_roundtrip_docx(self):
        from docx import Document
        r = _run([os.path.join(self.tmp, "informe_anon.docx"), "--restaurar"])
        self.assertEqual(r.returncode, 0, r.stderr)
        out = " ".join(p.text for p in Document(os.path.join(self.tmp, "informe_anon_restaurado.docx")).paragraphs)
        self.assertIn("María López", out)
        self.assertIn("12345678Z", out)


if __name__ == "__main__":
    unittest.main(verbosity=2)
